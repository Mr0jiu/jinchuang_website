from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path("outputs") / "网站数据核对与修正版报告.docx"

def set_font(run, name="Calibri", size=11, color="000000", bold=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_text(cell, text, bold=False, color="000000"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(str(text))
    set_font(r, size=9.5, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(int(width * 1440)))
            tcW.set(qn("w:type"), "dxa")

def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color="1F3A5F")
        shade(table.rows[0].cells[i], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table

def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size={1:16,2:13,3:12}[level], color="2E74B5" if level < 3 else "1F4D78", bold=True)
    return p

def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix), bold=True)
        set_font(p.add_run(text[len(bold_prefix):]))
    else:
        set_font(p.add_run(text))
    return p

def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for level, size, color, before, after in [(1,16,"2E74B5",16,8),(2,13,"2E74B5",12,6),(3,12,"1F4D78",8,4)]:
        st = styles[f"Heading {level}"]
        st.font.name = "Calibri"; st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size); st.font.color.rgb = RGBColor.from_string(color); st.font.bold = True
        st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)
    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("网站数据核对报告 · 2026-08-27"), size=9, color="6B7280")
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("数据核对留档"), size=9, color="6B7280")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("网站数据核对与修正版报告")
    set_font(r, size=23, color="0B2545", bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    set_font(p.add_run("页面数据与 D:\\GitHub\\jinchuang 源数据的对照、问题说明及修正结果"), size=13, color="555555")
    for label, value in [("核对日期", "2026-08-27"), ("核对对象", "D:\\GitHub\\jinchuang_website_code 页面 / D:\\GitHub\\jinchuang 数据目录与 outputs/mvp"), ("修正状态", "已修正首页统计口径并刷新数据库缓存")]:
        add_body(doc, f"{label}：{value}", label + "：")

    add_heading(doc, "一、结论摘要")
    add_body(doc, "页面贷款库与源影像目录已完成一一对应核对：源目录包含 3,254 个贷款目录、16,270 张影像，页面数据库同样包含 3,254 笔贷款，image_dir 无缺失、无多余。")
    add_body(doc, "发现并修正 1 项关键问题：旧首页缓存使用了历史 FAISS 面签清单中的 4,215 条记录，而当前页面实际贷款库只有 3,254 条面签特征，导致首页面签数、配对数及高相似统计被放大。现已改为以页面数据库 loans.face_feature 为唯一统计口径。")
    add_body(doc, "本报告中的“正确口径”指当前页面数据库与源影像目录一致后的统计结果；模型评估/监控报告仍属于独立产物口径，不与首页业务总览数字混用。")

    add_heading(doc, "二、源数据与页面数据对照")
    add_table(doc, ["核对项", "源数据", "页面修正版", "结果"], [
        ("贷款目录 / 贷款笔数", "3,254 个 loan_* 目录", "3,254 笔 loans", "一致"),
        ("影像总数", "16,270 张；每笔 5 张", "16,270 张（贷款数 × 5）", "一致"),
        ("面签影像", "3,254 张 face_signing.jpg", "3,254 条 face_feature", "一致"),
        ("影像清单状态", "data_manifest.csv：16,270 条均 ok", "页面目录映射无缺失/多余", "一致"),
        ("客户数", "页面业务库生成的客户主数据", "2,517 位", "页面口径"),
    ], [2.0, 1.65, 1.65, 1.2])
    add_body(doc, "源文件依据：D:\\GitHub\\jinchuang\\tmp\\v4_release\\extracted\\data、D:\\GitHub\\jinchuang\\outputs\\mvp\\data_manifest.csv；页面依据：D:\\GitHub\\jinchuang_website_code\\web1\\data.db。")

    add_heading(doc, "三、修正后的首页正确数据")
    add_table(doc, ["指标", "正确值", "说明"], [
        ("客户数", "2,517", "customers 表"),
        ("贷款数", "3,254", "loans 表"),
        ("影像总数", "16,270", "3,254 × 5"),
        ("面签影像数", "3,254", "当前 loans.face_feature 全量"),
        ("总配对数", "5,292,631", "3,254 × 3,253 ÷ 2"),
        ("高相似面签数（≥97%）", "2,336", "影像级去重口径"),
        ("高相似率", "71.79%", "2,336 ÷ 3,254"),
        ("同客户高相似", "717", "同客户优先归类"),
        ("跨客户高相似", "1,619", "排除同客户后的互斥值"),
        ("涉及贷款", "2,336", "至少命中一条高相似关系的贷款"),
        ("待复核", "882", "verify_status = F"),
        ("多笔授信客户 / 单笔客户", "884 / 2,370", "客户授信结构"),
    ], [2.2, 1.45, 2.85])

    add_heading(doc, "四、业务类型与行为分布")
    add_table(doc, ["业务类型", "贷款数", "同客户高相似", "跨客户高相似"], [
        ("商户易贷", "1,078", "255", "533"),
        ("锡微贷", "1,098", "219", "573"),
        ("消费贷", "1,078", "243", "513"),
    ], [2.2, 1.2, 1.55, 1.55])
    add_body(doc, "贷款行为：同客户复用 717、跨客户复用 1,619、正常 918；其中正常贷款里的多笔授信为 115 笔。")

    add_heading(doc, "五、模型指标逐项核对")
    add_body(doc, "网页报告页与实验页现统一使用 D:\\GitHub\\jinchuang_website_code\\jinchuang_v4\\code\\outputs\\mvp 这份 3,254 条面签数据对应的正式产物。旧的 D:\\GitHub\\jinchuang\\outputs\\mvp 仍保留 4,215 条历史清单，不再作为当前网页指标来源。")
    add_table(doc, ["指标组", "指标", "网页正确值", "核对来源"], [
        ("五类影像分类", "Accuracy / Macro-F1", "100.00% / 100.00%", "classification_metrics.json test"),
        ("分类测试集", "样本量", "2,440（5 类 × 488）", "classification_metrics.json test"),
        ("Stage1 配对级", "Precision / Recall / F1", "99.75% / 99.62% / 99.68%", "two_stage_summary.json"),
        ("Stage1 配对级", "ROC-AUC / 阈值", "100.00% / 0.94", "two_stage_summary.json"),
        ("Stage1 组级", "Precision / Recall / F1", "98.23% / 95.24% / 96.71%", "two_stage_summary.json"),
        ("Stage1 组级", "ROC-AUC / 阈值", "99.65% / 0.39", "two_stage_summary.json"),
        ("最终报告", "Precision / Recall / F1 / ROC-AUC", "100.00% / 100.00% / 100.00% / 100.00%", "final_report_metrics_after_review_labels"),
        ("Stage1 规模", "唯一无向配对 / 预测相似", "25,001 / 3,166", "two_stage_summary.json"),
        ("Stage2", "跨客户套用 / 同客户复用 / 同客户多笔授信", "2,073 / 83 / 1,010", "stage2_fraud_type_report.csv"),
        ("风险监控", "总配对 / 可疑配对 / 风险聚类", "25,001 / 11,447 / 247", "fraud_monitoring_summary.json"),
    ], [1.35, 1.65, 2.0, 1.5])
    add_body(doc, "页面显示百分比由原始小数 × 100 后保留两位；模型原始值已在上述正式 JSON 文件中保留。分类器的 100% 指当前测试集结果，不代表对未知业务数据的必然准确率。")

    add_heading(doc, "六、代码修正内容")
    add_body(doc, "已修改 dashboard_ui/refresh_stats.py：刷新首页缓存时不再优先读取可能包含历史补充样本的 FAISS manifest，而是直接读取当前页面数据库 loans.face_feature，并以 image_dir 建立客户关系映射。")
    add_body(doc, "已执行刷新：python dashboard_ui/refresh_stats.py；已通过页面接口 /api/stats 复核，返回 face_images=3,254、total_pairs=5,292,631、high_similar_pairs=2,336。")

    add_heading(doc, "七、留档建议")
    add_body(doc, "以后新增或替换贷款影像后，应先更新 loans.face_feature，再点击页面“刷新数据”；不要直接以历史 FAISS 清单条数作为首页业务总览的面签数量。")
    add_body(doc, "本次核对未发现当前源影像目录与页面贷款目录之间的缺失或多余记录。")
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT.resolve())

if __name__ == "__main__":
    main()
