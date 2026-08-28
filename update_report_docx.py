from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

src = Path(r'D:\GitHub\jinchuang_website_code\jinchuang_v4\code\outputs\mvp\mvp_complete_technical_report.docx')
dst = Path(r'D:\GitHub\jinchuang_website_code\jinchuang_v4\code\outputs\mvp\mvp_complete_technical_report_updated.docx')
doc = Document(src)

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_text(cell, text, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(9)

anchor = next((p for p in doc.paragraphs if p.text.strip() == '8. 网页数据一致性与缺口'), None)
if anchor is None:
    raise RuntimeError('anchor paragraph not found')

heading = OxmlElement('w:p')
anchor._p.addprevious(heading)
hp = anchor._element.getprevious()
from docx.text.paragraph import Paragraph
heading_para = Paragraph(hp, anchor._parent)
heading_para.style = doc.styles['Heading 2']
heading_para.add_run('阈值实验明细（与页面阈值档位一致）')

note = OxmlElement('w:p')
anchor._p.addprevious(note)
note_para = Paragraph(anchor._element.getprevious(), anchor._parent)
note_para.style = doc.styles['Normal']
note_para.add_run('以下结果来自当前 outputs/mvp 的 FAISS Top-5 全量面签离线回放。页面首页仍展示影像级命中数；本表展示阈值评估指标，二者统计对象不同。')

rows = [
    ('94.0%', '79.93%', '98.90%', '88.41%', '2,825'),
    ('94.5%', '82.07%', '98.47%', '89.53%', '2,739'),
    ('95.0%', '84.76%', '97.94%', '90.88%', '2,638'),
    ('95.5%', '88.09%', '97.20%', '92.42%', '2,519'),
    ('96.0%', '90.88%', '96.06%', '93.40%', '2,413'),
    ('96.5%', '93.35%', '94.66%', '94.00%', '2,315'),
    ('97.0%', '95.75%', '92.86%', '94.29%', '2,214'),
    ('97.5%', '97.84%', '91.11%', '94.35%', '2,126'),
    ('98.0%', '99.02%', '88.83%', '93.65%', '2,048'),
]
tbl = doc.add_table(rows=1, cols=5)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl.style = 'Table Grid'
for i, h in enumerate(['阈值', '精确率', '召回率', 'F1', '需复核/命中数']):
    set_cell_text(tbl.rows[0].cells[i], h, True)
    set_cell_shading(tbl.rows[0].cells[i], 'D9EAF7')
for row in rows:
    cells = tbl.add_row().cells
    for i, value in enumerate(row):
        set_cell_text(cells[i], value)
        cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
for row in tbl.rows:
    for cell in row.cells:
        cell.width = Inches(1.25)

tbl_el = tbl._element
anchor._p.addprevious(tbl_el)

doc.save(dst)
print(dst)
