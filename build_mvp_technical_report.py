from __future__ import annotations

import json
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "jinchuang_v4" / "code" / "outputs" / "mvp"
ASSETS = OUT / "technical_report_assets"
ASSETS.mkdir(exist_ok=True)


def load(name):
    return json.loads((OUT / name).read_text(encoding="utf-8"))


run = load("run_summary.json")
cls = load("classification_metrics.json")
ts = load("two_stage_summary.json")
fm = load("fraud_monitoring_summary.json")
stage1 = ts["stage1"]
pair = stage1["pair_level_split"]["metrics"]
group = stage1["group_level_split"]["metrics"]

BLUE = "#1263e6"
RED = "#ef4655"
ORANGE = "#f59a35"
GREEN = "#12a675"


FONT_CANDIDATES = ["C:/Windows/Fonts/msyh.ttc", "C:/Windows/Fonts/simhei.ttf", "C:/Windows/Fonts/arial.ttf"]
FONT = next((x for x in FONT_CANDIDATES if Path(x).exists()), None)
def font(size, bold=False):
    if FONT:
        return ImageFont.truetype(FONT, size)
    return ImageFont.load_default()

def canvas(title, width=1200, height=560):
    im = Image.new("RGB", (width, height), "white")
    d = ImageDraw.Draw(im); d.text((50, 28), title, fill="#0B1F44", font=font(28, True))
    return im, d

def savefig(im, name):
    path = ASSETS / name; im.save(path); return path


def chart_data_quality():
    im, d = canvas("数据质量检查：有效图片率 100%")
    vals = [run["total_images"], run["valid_images"], run["bad_images"]]
    colors = [BLUE, GREEN, RED]; labels = ["总图片", "有效图片", "异常图片"]
    base, top, bw = 450, 120, 180; scale = 300 / max(vals)
    for i, (v, label, color) in enumerate(zip(vals, labels, colors)):
        x = 160 + i * 310; y = base - int(v * scale)
        d.rectangle((x, y, x+bw, base), fill=color); d.text((x+20, base+18), label, fill="#405a78", font=font(21)); d.text((x+35, y-32), f"{v:,}", fill="#0B1F44", font=font(22, True))
    return savefig(im, "01_data_quality.png")


def chart_classification():
    im, d = canvas("五类影像分类准确率"); labels = ["训练", "验证", "测试"]
    vals = [cls[x]["accuracy"] * 100 for x in ["train", "val", "test"]]
    for i, (v, label, color) in enumerate(zip(vals, labels, ["#8bb8f7", "#4d91ef", BLUE])):
        x=240+i*300; y=450-int((v-95)*55); d.rectangle((x,y,x+150,450),fill=color); d.text((x+35,465),label,fill="#405a78",font=font(21)); d.text((x+30,y-32),f"{v:.2f}%",fill="#0B1F44",font=font(22,True))
    d.text((60, 450), "95%", fill="#8291a8", font=font(17)); d.text((60, 120), "101%", fill="#8291a8", font=font(17))
    return savefig(im, "02_classification.png")


def chart_stage1():
    im, d = canvas("Stage1 配对级与组级评估对比", height=620)
    labels = ["Precision", "Recall", "F1", "ROC-AUC"]
    pair_vals = [pair[x] * 100 for x in ["precision", "recall", "f1", "roc_auc"]]
    group_vals = [group[x] * 100 for x in ["precision", "recall", "f1", "roc_auc"]]
    for i, label in enumerate(labels):
        x=130+i*260; py=500-int((pair_vals[i]-90)*35); gy=500-int((group_vals[i]-90)*35); d.rectangle((x,py,x+65,500),fill=BLUE); d.rectangle((x+80,gy,x+145,500),fill=ORANGE); d.text((x,520),label,fill="#405a78",font=font(18)); d.text((x,py-28),f"{pair_vals[i]:.2f}",fill="#0B1F44",font=font(14)); d.text((x+78,gy-28),f"{group_vals[i]:.2f}",fill="#0B1F44",font=font(14))
    d.rectangle((850,80,875,105),fill=BLUE); d.text((885,80),"配对级",fill="#405a78",font=font(18)); d.rectangle((850,125,875,150),fill=ORANGE); d.text((885,125),"组级",fill="#405a78",font=font(18))
    return savefig(im, "03_stage1_metrics.png")


def chart_stage2():
    im, d = canvas("Stage2 全量相似结果类型分布（共 3,166 对）", height=580)
    labels = ["跨客户风险", "同客户复用", "正常续贷/同客户变化"]
    vals = [2073, 83, 1010]
    for i,(label,v,color) in enumerate(zip(labels,vals,[RED,ORANGE,GREEN])):
        y=150+i*110; d.text((70,y+15),label,fill="#405a78",font=font(21)); d.rectangle((390,y,390+int(v/2073*650),y+48),fill=color); d.text((410+int(v/2073*650),y+10),f"{v:,}",fill="#0B1F44",font=font(20,True))
    return savefig(im, "04_stage2_types.png")


def chart_fraud():
    im, d = canvas("欺诈监控：可疑配对与风险等级", height=580)
    d.ellipse((100,150,430,480), fill="#dbe8f8"); ratio=fm["suspicious_pairs"]/fm["total_pairs"]; d.pieslice((100,150,430,480), start=90, end=90-360*ratio, fill=RED); d.text((175,285),f"{ratio:.2%}",fill="#0B1F44",font=font(28,True)); d.text((185,500),"可疑配对占比",fill="#405a78",font=font(20))
    levels = ["critical", "high", "medium", "low"]
    vals = [fm["by_fraud_score_level"][x] for x in levels]
    for i,(level,v,color) in enumerate(zip(levels,vals,[RED,ORANGE,"#6ca3ed",GREEN])):
        x=560+i*140; h=int(v/11242*300); d.rectangle((x,470-h,x+70,470),fill=color); d.text((x+5,485),level,fill="#405a78",font=font(15)); d.text((x,450-h),f"{v:,}",fill="#0B1F44",font=font(14))
    d.text((560,90),"风险等级分布",fill="#0B1F44",font=font(23,True))
    return savefig(im, "05_fraud_monitoring.png")


charts = [chart_data_quality(), chart_classification(), chart_stage1(), chart_stage2(), chart_fraud()]


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tcPr.append(shd)


def set_cell(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text)); r.bold = bold; r.font.size = Pt(9)
    if color: r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers): set_cell(t.rows[0].cells[i], h, True, "FFFFFF"); shade(t.rows[0].cells[i], "1F4D78")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row): set_cell(cells[i], v)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths): row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet"); p.add_run(text); return p


doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Inches(0.75)
sec.left_margin = sec.right_margin = Inches(0.8)
styles = doc.styles
styles["Normal"].font.name = "Calibri"; styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei"); styles["Normal"].font.size = Pt(10.5)
for name, size, color in [("Title", 25, "0B1F44"), ("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 11, "1F4D78")]:
    s = styles[name]; s.font.name = "Calibri"; s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei"); s.font.size = Pt(size); s.font.color.rgb = RGBColor.from_string(color)

p = doc.add_paragraph(style="Title"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run("MVP 多模态金融影像智能相似度检测\n完整技术报告")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run("基于 outputs/mvp 正式产物 · 2026-08-27").italic = True
doc.add_paragraph("本报告对当前 MVP 的数据质量、材料分类、Stage1 相似检测、Stage2 风险解释、欺诈监控、页面数据口径及上线风险进行统一分析。所有数字均来自当前目录中的 JSON/CSV 汇总产物；模型结果、风险线索和人工确认案件严格区分。")

doc.add_heading("1. 执行摘要", level=1)
table(doc, ["核心指标", "结果", "解释"], [
    ["数据规模", "16,270 张图片 / 3,254 笔贷款", "5 类材料，每类 3,254 张"],
    ["数据有效率", "100.00%", "16,270 张全部通过检查"],
    ["分类测试准确率", "100.00%", "五类影像，测试集 2,440 张"],
    ["Stage1 组级 F1", "96.71%", "更接近新客户/新相似组泛化"],
    ["可疑配对", "11,447 / 25,001（45.79%）", "模型风险线索，不等于已确认欺诈"],
    ["运行耗时", "17.8 秒", "run_summary.json 记录的本次运行耗时"],
], [1.7, 1.7, 3.1])
doc.add_picture(str(charts[0]), width=Inches(6.2)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("图 1  数据质量检查结果。", style="Caption")

doc.add_heading("2. 数据与产物清单", level=1)
doc.add_paragraph("当前正式产物包含分类指标、运行汇总、两阶段汇总、欺诈监控汇总、人脸清单、分类器和 FAISS 索引。face_manifest.csv 共 3,254 条，状态全部为 ok，预测类别全部为 face_signing，置信度均值 0.997927，中位数 0.998370，最低值 0.970100。")
table(doc, ["文件", "用途", "状态"], [
    ["classification_metrics.json", "五类材料分类评估", "可用"],
    ["run_summary.json", "数据规模、阈值、耗时和运行元数据", "可用"],
    ["two_stage_summary.json", "Stage1/Stage2 评估与类型汇总", "可用"],
    ["fraud_monitoring_summary.json", "欺诈监控和风险图汇总", "可用"],
    ["face_manifest.csv", "面签影像索引与分类置信度", "可用"],
    ["stage1_similarity_report.csv 等明细", "样例、阈值曲线和风险最高样例", "当前目录未提供"],
], [2.5, 2.9, 1.1])

doc.add_heading("3. 五类影像分类", level=1)
doc.add_paragraph("模型为 google/siglip2-base-patch16-224，运行设备为 CUDA。训练、验证和测试集的五类 precision、recall、F1 均为 1.0000，混淆矩阵均为对角矩阵。该结果说明当前测试样本没有分类错误，但不能推断未来未知影像也能达到 100%。")
doc.add_picture(str(charts[1]), width=Inches(6.2)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("图 2  五类影像分类准确率。", style="Caption")

doc.add_heading("4. Stage1 相似检测", level=1)
doc.add_paragraph("Stage1 采用 recall-first 策略，目标召回率 95%，允许误报进入人工复核。原始方向性配对 32,540 行，去除 7,539 条对称重复记录后得到 25,001 个唯一无向配对，无自配对记录。")
table(doc, ["评估口径", "阈值", "Precision", "Recall", "F1", "ROC-AUC"], [
    ["配对级测试", "0.94", "99.75%", "99.62%", "99.68%", "99.9967%"],
    ["组级测试", "0.39", "98.23%", "95.24%", "96.71%", "99.6540%"],
], [1.4, .8, 1.0, .9, .8, 1.1])
doc.add_picture(str(charts[2]), width=Inches(6.2)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("图 3  Stage1 两种切分口径对比。组级结果更保守，应作为上线容量规划的主要依据。", style="Caption")
add_bullet(doc, "配对级切分可能共享相似组信息，指标偏乐观；组级隔离更接近跨客户泛化场景。")
add_bullet(doc, "全量回溯结果为 TP 3,166、FP 0、FN 0、TN 21,835，但它依赖当前标签定义，不是独立人工验证结果。")

doc.add_heading("5. Stage2 风险解释", level=1)
doc.add_paragraph("Stage2 不是独立的二分类测试集，而是对 Stage1 相似结果进行业务解释。当前 3,166 个最终相似对被归入跨客户风险、同客户复用和正常续贷/同客户变化。另有 33 对处于高相似但身份待确认状态。")
table(doc, ["类型", "数量", "占 3,166 对"], [["跨客户风险", "2,073", "65.48%"], ["同客户复用复核", "83", "2.62%"], ["正常续贷/同客户变化", "1,010", "31.90%"]], [2.4, 1.2, 1.6])
doc.add_picture(str(charts[3]), width=Inches(6.2)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("图 4  Stage2 全量相似结果类型分布。", style="Caption")

doc.add_heading("6. 欺诈监控与风险图", level=1)
doc.add_paragraph("风险监控层共分析 25,001 个配对，标记 11,447 个可疑配对，占 45.79%。其中跨客户欺诈 10,320 个，占可疑配对 90.15%；critical 级 11,242 个，占可疑配对 98.21%。这些数字表示需要关注的模型线索，不是已经人工确认的案件数。")
table(doc, ["监控指标", "数量", "占可疑配对"], [["same_customer_repeat", "1,054", "9.21%"], ["cross_customer_fraud", "10,320", "90.15%"], ["cross_customer_candidate", "73", "0.64%"], ["urgent", "3,768", "32.92%"], ["风险聚类", "247", "—"], ["最大风险簇", "1,506 个节点", "风险图节点的 56.47%"]], [2.6, 1.4, 1.5])
doc.add_picture(str(charts[4]), width=Inches(6.2)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("图 5  欺诈监控可疑占比与风险等级分布。", style="Caption")

doc.add_heading("7. 阈值与业务策略", level=1)
table(doc, ["策略参数", "当前值", "含义"], [["high_risk_threshold", "0.97", "业务高风险阈值"], ["medium_risk_threshold", "0.93", "中风险/预警阈值"], ["cross_customer_threshold", "0.95", "跨客户关系阈值"], ["same_customer_threshold", "0.92", "同客户关系阈值"], ["Stage1 配对级报告阈值", "0.94", "相似报告口径"], ["Stage1 组级测试阈值", "0.39", "组级泛化评估口径"]], [2.6, 1.1, 3.0])
doc.add_paragraph("业务策略阈值与模型验证阈值用途不同，页面应始终标注口径，不能把 0.97 直接描述为模型最优验证阈值。")

doc.add_heading("8. 网页数据一致性与缺口", level=1)
doc.add_paragraph("网页当前采用混合数据架构：首页统计和贷款详情来自 SQLite；技术报告和模型指标来自 outputs/mvp；智能检测使用当前模型桥、FAISS 和影像目录；操作记录来自 SQLite。该架构能保留业务实时状态，但数据库与 MVP 产物更新不同步时，首页数量可能与技术报告不同。")
add_bullet(doc, "已修正技术报告页旧的 99.87%、99.64%、99.76% 等初始展示值，并明确分类/配对级口径。")
add_bullet(doc, "已修正错误的 14,769 个影像对文案为 25,001 个唯一无向影像对。")
add_bullet(doc, "当前目录缺少 stage1_similarity_report.csv、stage2_fraud_type_report.csv、threshold_experiment.csv、fraud_monitoring.csv，因此网页无法恢复明细样例和阈值曲线。")

doc.add_heading("9. 结论与上线建议", level=1)
for text in [
    "将组级 F1 96.71%、Recall 95.24%作为更保守的上线性能基准；配对级指标作为补充。",
    "对分类 100% 结果增加真实业务噪声、跨来源样本和人工盲测，排查模板偏差或数据泄漏。",
    "11,447 个可疑配对全部进入人工队列可能造成较大审核压力，建议按风险簇、客户和业务金额聚合去重。",
    "优先核验 73 对身份待确认记录，并对最大风险簇做人工抽样和误报回流。",
    "补齐四类明细 CSV，使网页能展示阈值实验曲线、Stage1 样例和风险最高样例。",
    "所有接口返回 source、updated_at 和数据版本；前端只展示后端结果，不手写业务数字。",
]: add_bullet(doc, text)

doc.add_heading("10. 数据来源", level=1)
for x in ["classification_metrics.json", "run_summary.json", "two_stage_summary.json", "fraud_monitoring_summary.json", "face_manifest.csv", "dashboard_ui/index.html", "dashboard_ui/modern_server.py", "web1/model_bridge.py"]: add_bullet(doc, x)

footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER; footer.add_run("MVP 技术报告 · 千影一鉴团队 · 2026-08-27").font.size = Pt(8)
path = OUT / "mvp_complete_technical_report.docx"
doc.save(path)
print(path)
